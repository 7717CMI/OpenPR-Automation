from flask import Flask, render_template, request, send_file, redirect, url_for, flash, make_response, jsonify, session
import pandas as pd
import os
import subprocess
import time
import pyautogui
from functools import wraps
from werkzeug.utils import secure_filename
import win32com.client
from docx import Document
from datetime import datetime, timedelta
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import Select
from webdriver_manager.chrome import ChromeDriverManager
import threading
import logging
import re
import json
import requests

app = Flask(__name__)
app.secret_key = 'your_secret_key_change_this_in_production'
app.config['UPLOAD_FOLDER'] = 'uploads/'
app.config['MAX_CONTENT_LENGTH'] = 25 * 1024 * 1024  # 25MB max file size

# Hardcoded download filename and path
app.config['DOWNLOAD_FILENAME'] = 'ROB.xlsx'
app.config['DOWNLOAD_PATH'] = r'C:\Users\akshat\Desktop\RPA\\' + app.config['DOWNLOAD_FILENAME']

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Set up logging to capture output
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Add a global variable to track processing status
processing_status = {
    'active': False,
    'message': 'Ready',
    'progress': 0,
    'total': 0,
    'current_file': '',
    'logs': []
}

# GOOGLE TRENDS CONFIG
GOOGLE_TRENDS_CONFIG = {
    'SERP_API_KEY': '105e00ebbd95f55724ddecb6ed5bb32ecd068d9ba798c6f99b4ce0d17baebc58',
    'INTEREST_THRESHOLD': 50,
    'DAYS_ABOVE_THRESHOLD': 2,
    'TARGET_COUNTRIES': ['india', 'usa', 'china', 'germany', 'south korea', 'france', 'uk'],
    'TERMS_TO_REMOVE': ['market', 'size', 'analysis', 'report', 'industry', 'global'],
    'REQUEST_DELAY': 2,
    'REQUIRE_COUNTRIES_CHECK': False,
    'LENIENT_MODE': True
}

def allowed_file(filename):
    """Check if file extension is allowed"""
    ALLOWED_EXTENSIONS = {'xlsx', 'xls', 'csv'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def allowed_backend_file(filename):
    """Check if backend file extension is allowed"""
    BACKEND_EXTENSIONS = {'xlsx', 'xls'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in BACKEND_EXTENSIONS

def log_to_status(message):
    """Add a message to the processing status logs"""
    global processing_status
    processing_status['logs'].append(f"{datetime.now().strftime('%H:%M:%S')}: {message}")
    print(f"[LOG] {message}")

# ============================================================================
# HOME ROUTE
# ============================================================================

@app.route('/')
def index():
    return render_template('index.html')

# ============================================================================
# DOCUMENT PROCESSING ROUTES
# ============================================================================

@app.route('/document_processing', methods=['GET', 'POST'])
def document_processing():
    global processing_status
    
    if request.method == 'POST':
        try:
            # Get form data - use session data as defaults if available
            article_code = request.form.get('article_code') or request.form.get('open_pr_id') or session.get('open_pr_id', '6HA-2025-M6K439')
            author_name = request.form.get('author_name') or session.get('username', 'Vishwas tiwari')
            author_email = request.form.get('author_email') or session.get('email', 'vishwas@coherentmarketinsights.com')
            company_name = request.form.get('company_name', 'Coherent Market Insights')
            phone_number = request.form.get('phone_number') or session.get('mobile', '1234567890')
            
            # Power Automate output folder path
            custom_folder = request.form.get('custom_folder')
            if custom_folder:
                folder_path = custom_folder
            else:
                today = datetime.today()
                folder_path = rf'C:\Users\akshat\Desktop\RPA\Files\{today.year}\{today.strftime("%m")}\{today.strftime("%d")}'
            
            processing_mode = request.form.get('processing_mode', 'manual')
            
            # Validate paths before processing
            excel_path = r'C:\Users\akshat\Desktop\RPA\ROB.xlsx'
            
            # Check if required files exist
            validation_errors = []
            if not os.path.exists(excel_path):
                validation_errors.append(f"Excel file not found: {excel_path}")
            if not os.path.exists(folder_path):
                validation_errors.append(f"Folder not found: {folder_path}")
            
            if validation_errors:
                for error in validation_errors:
                    flash(error)
                return render_template('document_processing.html', 
                                     session_data={
                                         'username': session.get('username', ''),
                                         'email': session.get('email', ''),
                                         'mobile': session.get('mobile', ''),
                                         'open_pr_id': session.get('open_pr_id', '')
                                     })
            
            # Reset processing status
            processing_status = {
                'active': True,
                'message': 'Starting processing...',
                'progress': 0,
                'total': 0,
                'current_file': '',
                'logs': []
            }
            
            # Start processing in background thread
            if processing_mode == 'auto':
                threading.Thread(target=process_documents_auto_with_feedback, 
                               args=(folder_path, article_code, author_name, author_email, 
                                    company_name, phone_number)).start()
            else:
                threading.Thread(target=process_documents_manual_with_feedback, 
                               args=(folder_path, article_code, author_name, author_email, 
                                    company_name, phone_number)).start()
            
            flash('Processing started! Check the status page for updates.')
            return redirect(url_for('processing_status'))
            
        except Exception as e:
            flash(f'Error starting processing: {str(e)}')
            logger.error(f"Error in document_processing: {e}")
            return render_template('document_processing.html', 
                                 session_data={
                                     'username': session.get('username', ''),
                                     'email': session.get('email', ''),
                                     'mobile': session.get('mobile', ''),
                                     'open_pr_id': session.get('open_pr_id', '')
                                 })
    
    # Pre-populate form with session data if available
    return render_template('document_processing.html', 
                         session_data={
                             'username': session.get('username', ''),
                             'email': session.get('email', ''),
                             'mobile': session.get('mobile', ''),
                             'open_pr_id': session.get('open_pr_id', '')
                         })

@app.route('/processing_status')
def processing_status_page():
    return render_template('processing_status.html')

@app.route('/api/get_processing_status')
def get_processing_status():
    """API endpoint to get current processing status"""
    global processing_status
    return jsonify(processing_status)

# ============================================================================
# DOCUMENT PROCESSING FUNCTIONS
# ============================================================================

def convert_doc_to_docx(doc_path, output_path=None):
    """Convert .doc file to .docx format"""
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(doc_path)
        if not output_path:
            output_path = os.path.splitext(doc_path)[0] + ".docx"
        doc.SaveAs(output_path, FileFormat=16)
        doc.Close()
        word.Quit()
        return output_path
    except Exception as e:
        log_to_status(f"Error converting doc to docx: {e}")
        return None

def text_of_press_release(doc_path, start_index=21, end_index=-8):
    # Load the DOCX file
    doc = Document(doc_path)

    # Extract only V4 section paragraphs
    v4_paragraphs = []
    v4_found = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if text == "V4":
            v4_found = True
            continue
        elif text.startswith("Version 5") and v4_found:
            break
        elif v4_found:
            v4_paragraphs.append(para)

    # If V4 not found, fall back to all paragraphs
    if not v4_paragraphs:
        v4_paragraphs = doc.paragraphs

    # Extract text with formatting preservation
    formatted_lines = []
    for para in v4_paragraphs:
        text = para.text.strip()
        if not text or text.replace('_', '').replace('-', '').strip() == "":
            if formatted_lines and formatted_lines[-1] != "":
                formatted_lines.append("")
            continue
        formatted_lines.append(text)

    saved = "\n".join(formatted_lines)
    words = saved.split()

    if len(words) > abs(end_index):
        chunk = " ".join(words[start_index:end_index])

        # Add line breaks before section headers
        section_headers = [
            'Market Size and Overview',
            'Key Takeaways',
            'Segments Covered:',
            'Growth Factors',
            'Market Trends',
            'Actionable Insights',
            'Key Players',
            'FAQs'
        ]
        for header in section_headers:
            chunk = chunk.replace(header, f"\n\n{header}")

        # Remove dashes before content
        chunk = re.sub(r'-{2,}', '', chunk)

        # Add line breaks before bullet points
        chunk = chunk.replace(' - ', '\n- ')

        # Add line breaks before FAQ labels
        chunk = re.sub(r'\s*(FAQ?s?:?)', r'\n\n\1\n\n', chunk)

        # Ensure each numbered FAQ starts on a new line
        chunk = re.sub(r'\s*(\d+\.\s)', r'\n\1', chunk)

        # Add proper spacing around phrase + link combinations using regex
        patterns = [
            r"(Explore the Entire Market Report here:\s*)(https://www\.coherentmarketinsights\.com/market-insight/[^\s]+)",
            r"(Request for Sample Copy of the Report here\s*:\s*)(https://www\.coherentmarketinsights\.com/insight/request-sample/[^\s]+)",
            r"(Get Instant Access! Purchase Research Report and Receive a 25% Discount:\s*)(https://www\.coherentmarketinsights\.com/insight/buy-now/[^\s]+)"
        ]
        
        for pattern in patterns:
            chunk = re.sub(pattern, r"\n\n\1\2\n", chunk)
        
        chunk = re.sub(r'\n{3,}', '\n\n', chunk)
        chunk = chunk.strip()
        return chunk
    else:
        return "Text not found."
    
def run_selenium_automation(article_code, article_title, multiline_text, author_name, 
                          author_email, company_name, phone_number):
    """Run Selenium automation for press release submission"""
    try:
        log_to_status("Starting Selenium automation...")
        chromedriver_path = ChromeDriverManager().install()
        options = Options()
        options.add_argument("--start-maximized")
        
        cService = Service(executable_path=chromedriver_path)
        driver = webdriver.Chrome(service=cService, options=options)
        driver.get('https://www.openpr.com/')
        
        # Handle cookie consent
        try:
            reject = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.XPATH, '//*[@id="cmpbntnotxt"]'))
            )
            reject.click()
        except:
            pass
        
        # Navigate to submit page
        submit = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="navbarText"]/ul/li[3]/a'))
        )
        submit.click()
        
        # Enter article code
        input_box = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="code"]'))
        )
        input_box.clear()
        input_box.send_keys(article_code)
        
        # Submit code
        try:
            submit2 = WebDriverWait(driver, 5).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, '#main > div > div > div:nth-child(5) > div > form > button'))
            )
            submit2.click()
        except:
            submit2 = WebDriverWait(driver, 5).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, '#main > div > div > div:nth-child(6) > div > form > button'))
            )
            submit2.click()
        
        # Fill form fields
        name = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[1]/div/input'))
        )
        name.send_keys(author_name)
        
        email = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[2]/div/input'))
        )
        email.clear()
        email.send_keys(author_email)
        
        pr_agency = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[3]/div/input'))
        )
        pr_agency.clear()
        pr_agency.send_keys(author_name)
        
        number = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[4]/div/input'))
        )
        number.clear()
        number.send_keys(phone_number)
        
        ComName = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="archivnmfield"]'))
        )
        ComName.clear()
        ComName.send_keys(company_name)
        
        s1 = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="popup-archiv"]/div/a[1]'))
        )
        s1.click()
        
        Category_element = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[6]/div/select'))
        )
        Select(Category_element).select_by_visible_text("Arts & Culture")
        
        title = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[7]/div/input'))
        )
        title.clear()
        title.send_keys(article_title)
        
        text = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="inhalt"]'))
        )
        text.clear()
        text.send_keys(multiline_text)
        
        about = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[9]/div/textarea'))
        )
        about.clear()
        multi = """Contact Us:
        Mr. Shah
        Coherent Market Insights Pvt. Ltd,
        U.S.: + 12524771362
        U.K.: +442039578553
        AUS: +61-2-4786-0457
        INDIA: +91-848-285-0837
        ✉ Email: sales@coherentmarketinsights.com
        About Us:
        Coherent Market Insights leads into data and analytics, audience measurement, consumer behaviors, and market trend analysis. From shorter dispatch to in-depth insights, CMI has exceled in offering research, analytics, and consumer-focused shifts for nearly a decade. With cutting-edge syndicated tools and custom-made research services, we empower businesses to move in the direction of growth. We are multifunctional in our work scope and have 450+ seasoned consultants, analysts, and researchers across 26+ industries spread out in 32+ countries.")
        """
        about.send_keys(multi)
        
        address = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[10]/div/textarea'))
        )
        address.clear()
        address.send_keys("123 Test Street, Test City, Test Country")
        
        image = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="bild"]'))
        )
        image.clear()
        image.send_keys(r"C:\Users\akshat\Desktop\code\Market Analysis 2025.jpg")
        
        caption = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[12]/div/input'))
        )
        caption.clear()
        caption.send_keys("This is a test caption for the image.")
        
        notes = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[13]/div/textarea'))
        )
        notes.clear()
        notes.send_keys("This is a test notes section for the press release submission.")
        
        # Agree to terms
        tick1 = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="input-agb"]'))
        )
        tick1.click()
        
        tick2 = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="input-ds"]'))
        )
        tick2.click()
        
        # Submit form
        final = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/button'))
        )
        final.click()
        
        time.sleep(5)
        log_to_status("Selenium automation completed successfully")
        return True
        
    except Exception as e:
        log_to_status(f"Selenium automation error: {e}")
        try:
            driver.quit()
        except:
            pass
        return False

def process_documents_auto_with_feedback(folder_path, article_code, author_name, author_email, company_name, phone_number):
    """Process documents automatically with status feedback"""
    global processing_status
    
    try:
        log_to_status(f"Starting auto processing. Folder: {folder_path}")
        
        excel_path = r'C:\Users\akshat\Desktop\RPA\ROB.xlsx'
        
        # Load Excel file
        log_to_status("Loading Excel file...")
        keywords_df = pd.read_excel(excel_path)
        market_names = keywords_df['Market Name'].dropna().tolist()
        
        processing_status['total'] = len(market_names)
        log_to_status(f"Found {len(market_names)} market names to process")
        
        processed_count = 0
        
        for i, market_name in enumerate(market_names):
            processing_status['progress'] = i
            processing_status['current_file'] = market_name
            processing_status['message'] = f"Auto-processing {i+1} of {len(market_names)}: {market_name}"
            
            doc_file = f"ROB_{market_name}.doc"
            doc_path = os.path.join(folder_path, doc_file)
            
            log_to_status(f"Looking for file: {doc_path}")
            
            if os.path.exists(doc_path):
                log_to_status(f"Processing: {market_name}")
                
                # Convert doc to docx
                processing_status['message'] = f"Converting {market_name} to DOCX..."
                docx_path = convert_doc_to_docx(doc_path)
                
                if not docx_path:
                    log_to_status(f"ERROR: Could not convert {doc_path} to docx")
                    continue
                
                # Extract text
                processing_status['message'] = f"Extracting text from {market_name}..."
                multiline_text = text_of_press_release(docx_path)
                article_title = f"{market_name} Size, Trends, and Growth Forecast 2025-2032"
                
                # Run automation
                processing_status['message'] = f"Submitting {market_name} via automation..."
                success = run_selenium_automation(article_code, article_title, multiline_text, 
                                                author_name, author_email, company_name, phone_number)
                
                if success:
                    log_to_status(f"SUCCESS: Published {market_name}")
                    processed_count += 1
                else:
                    log_to_status(f"FAILED: Could not publish {market_name}")
                
                time.sleep(10)
                
            else:
                log_to_status(f"ERROR: File not found: {doc_path}")
        
        processing_status['active'] = False
        processing_status['message'] = f"Auto-processing complete! Published {processed_count} of {len(market_names)} articles"
        log_to_status(f"Auto processing complete. Published {processed_count} articles.")
        
    except Exception as e:
        processing_status['active'] = False
        processing_status['message'] = f"Error: {str(e)}"
        log_to_status(f"EXCEPTION: Auto processing error: {e}")

def process_documents_manual_with_feedback(folder_path, article_code, author_name, author_email, company_name, phone_number):
    """Process documents with manual intervention and status feedback"""
    global processing_status
    
    try:
        log_to_status(f"Starting manual processing. Folder: {folder_path}")
        
        excel_path = r'C:\Users\akshat\Desktop\RPA\ROB.xlsx'
        
        # Load Excel file
        log_to_status("Loading Excel file...")
        keywords_df = pd.read_excel(excel_path)
        market_names = keywords_df['Market Name'].dropna().tolist()
        
        processing_status['total'] = len(market_names)
        log_to_status(f"Found {len(market_names)} market names to process")
        
        processed_count = 0
        
        for i, market_name in enumerate(market_names):
            processing_status['progress'] = i
            processing_status['current_file'] = market_name
            processing_status['message'] = f"Processing {i+1} of {len(market_names)}: {market_name}"
            
            doc_file = f"ROB_{market_name}.doc"
            doc_path = os.path.join(folder_path, doc_file)
            
            log_to_status(f"Looking for file: {doc_path}")
            
            if os.path.exists(doc_path):
                log_to_status(f"Processing: {market_name}")
                
                # Convert doc to docx
                processing_status['message'] = f"Converting {market_name} to DOCX..."
                docx_path = convert_doc_to_docx(doc_path)
                
                if not docx_path:
                    log_to_status(f"ERROR: Could not convert {doc_path} to docx")
                    continue
                
                # Extract text
                processing_status['message'] = f"Extracting text from {market_name}..."
                multiline_text = text_of_press_release(docx_path)
                article_title = f"{market_name} Market Insights"
                
                # Run automation
                processing_status['message'] = f"Submitting {market_name} via automation..."
                success = run_selenium_automation(article_code, article_title, multiline_text, 
                                                author_name, author_email, company_name, phone_number)
                
                if success:
                    log_to_status(f"Published {market_name}")
                    processed_count += 1
                
                time.sleep(5)
                
            else:
                log_to_status(f"ERROR: File not found: {doc_path}")
        
        processing_status['active'] = False
        processing_status['message'] = f"Processing complete! Published {processed_count} of {len(market_names)} articles"
        log_to_status(f"Manual processing complete. Published {processed_count} articles.")
        
    except Exception as e:
        processing_status['active'] = False
        processing_status['message'] = f"Error: {str(e)}"
        log_to_status(f"EXCEPTION: Manual processing error: {e}")

# ============================================================================
# ROB PROCESSING ROUTES
# ============================================================================

@app.route('/rob', methods=['GET', 'POST'])
def rob():
    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        open_pr_id = request.form.get('open_pr_id')
        mobile = request.form.get('mobile')
        extract_count = int(request.form.get('extract_count', 200))

        # Validate required fields
        if not all([username, email, open_pr_id, mobile, extract_count]):
            flash('All fields are required!')
            return redirect(request.url)

        file = request.files.get('file')
        if not file or file.filename == '':
            flash('Excel file is required!')
            return redirect(request.url)

        if not allowed_file(file.filename):
            flash('Only Excel files (.xlsx, .xls) and CSV files are allowed!')
            return redirect(request.url)

        # Use secure_filename to avoid path issues
        filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(input_path)

        # Store user data in session for later use
        session['username'] = username
        session['email'] = email
        session['open_pr_id'] = open_pr_id
        session['mobile'] = mobile
        
        return redirect(url_for('process_rob', file_path=input_path,
                                username=username, email=email,
                                open_pr_id=open_pr_id, mobile=mobile,
                                extract_count=extract_count))
    return render_template('rob.html')

@app.route('/process_rob')
def process_rob():
    file_path = request.args.get('file_path')
    username = request.args.get('username')
    email = request.args.get('email')
    open_pr_id = request.args.get('open_pr_id')
    mobile = request.args.get('mobile')
    extract_count = int(request.args.get('extract_count', 200))

    if not file_path or not os.path.exists(file_path):
        flash('Missing or invalid file path')
        return redirect(url_for('rob'))

    try:
        # Read the cleaned ROB file
        if file_path.endswith('.csv'):
            df_original = pd.read_csv(file_path)
        else:
            df_original = pd.read_excel(file_path, engine='openpyxl')

        total_rows = len(df_original)
        
        if total_rows < extract_count:
            flash(f'⚠️ File only has {total_rows} rows, but you requested {extract_count} rows!')
            extract_count = total_rows

        # Step 1: Extract top N rows
        extracted_rows = df_original.head(extract_count).copy()
        
        # Step 2: Get remaining rows (original minus extracted)
        remaining_rows = df_original.iloc[extract_count:].copy()

        # Step 3: Create timestamp for remaining file
        today = datetime.today()
        timestamp = f"{today.year}_{today.month:02d}_{today.day:02d}"
        
        # Step 4: Save extracted rows as ROB.xlsx to Desktop/RPA
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        rpa_folder = os.path.join(desktop_path, "RPA")
        
        if not os.path.exists(rpa_folder):
            os.makedirs(rpa_folder)
        
        rob_output_path = os.path.join(rpa_folder, "ROB.xlsx")
        extracted_rows.to_excel(rob_output_path, index=False)
        
        # Step 5: Save remaining rows with timestamp
        remaining_filename = f"cleaned_rob_{timestamp}.xlsx"
        remaining_output_path = os.path.join(app.config['UPLOAD_FOLDER'], remaining_filename)
        remaining_rows.to_excel(remaining_output_path, index=False)
        
        # Step 6: Store info in session for the result page
        session['rob_file_path'] = rob_output_path
        session['remaining_file_path'] = remaining_output_path
        session['remaining_filename'] = remaining_filename
        session['extracted_count'] = extract_count
        session['remaining_count'] = len(remaining_rows)
        session['total_count'] = total_rows
        
        flash(f'✅ Successfully processed {total_rows} rows!')
        flash(f'📁 Remaining {len(remaining_rows)} rows → {remaining_filename} (ready for download)')
        
        # Use render_template instead of redirect
        return render_template('rob_result.html',
                             extracted_count=extract_count,
                             remaining_count=len(remaining_rows),
                             total_count=total_rows,
                             username=username,
                             records_processed=total_rows)

    except Exception as e:
        flash(f'❌ Error processing ROB file: {str(e)}')
        return redirect(url_for('rob'))

@app.route('/download_remaining_rob')
def download_remaining_rob():
    """Download the remaining ROB file (original minus extracted rows)"""
    try:
        remaining_file_path = session.get('remaining_file_path')
        remaining_filename = session.get('remaining_filename', 'cleaned_rob_remaining.xlsx')
        
        if remaining_file_path and os.path.exists(remaining_file_path):
            return send_file(remaining_file_path, as_attachment=True, download_name=remaining_filename)
        else:
            flash('❌ Remaining ROB file not found. Please process a file first.')
            return redirect(url_for('rob'))
    except Exception as e:
        flash(f'❌ Error downloading remaining file: {str(e)}')
        return redirect(url_for('rob'))
@app.route('/download_extracted_rob')
def download_extracted_rob():
    """Download the extracted ROB.xlsx file and trigger Power Automate"""
    try:
        rob_file_path = session.get('rob_file_path')
        
        if rob_file_path and os.path.exists(rob_file_path):
            # Set a flag in session to trigger Power Automate after download
            session['trigger_power_automate'] = True
            session['power_automate_triggered_at'] = time.time()
            
            # Start Power Automate in background thread with 5-second delay
            threading.Thread(target=delayed_power_automate_trigger, args=(5,)).start()
            
            return send_file(rob_file_path, as_attachment=True, download_name='ROB.xlsx')
        else:
            flash('❌ ROB.xlsx file not found. Please process a file first.')
            return redirect(url_for('rob'))
    except Exception as e:
        flash(f'❌ Error downloading ROB file: {str(e)}')
        return redirect(url_for('rob'))

def delayed_power_automate_trigger(delay_seconds=5):
    """Trigger Power Automate after a delay"""
    try:
        print(f"⏳ Waiting {delay_seconds} seconds before triggering Power Automate...")
        time.sleep(delay_seconds)
        
        print("🤖 Auto-triggering Power Automate after ROB download...")
        
        # Trigger Power Automate Desktop flow
        pad_exe_path = r"C:\Program Files (x86)\Power Automate Desktop\PAD.Console.Host.exe"
        flow_name = "Paid PR - Files Downloader"
        
        if os.path.exists(pad_exe_path):
            command = f'"{pad_exe_path}" -flow "{flow_name}"'
            
            result = subprocess.run(command, shell=True, check=True, text=True, capture_output=True)
            print(f"✅ Power Automate triggered successfully: {result.stdout}")
            
            # Wait for PAD to load
            time.sleep(5)
            
            # Click the flow button
            flow_button_coordinates = (463, 395)
            pyautogui.click(flow_button_coordinates)
            print("✅ Power Automate flow clicked and started")
            
        else:
            print("❌ Power Automate Desktop not found")
            
    except subprocess.CalledProcessError as e:
        print(f"❌ Error triggering Power Automate: {e.stderr}")
    except Exception as e:
        print(f"❌ Unexpected error in Power Automate trigger: {e}")

# Alternative API endpoint for manual triggering if needed
@app.route('/api/auto_trigger_power_automate', methods=['POST'])
def auto_trigger_power_automate():
    """API endpoint for auto-triggering Power Automate"""
    try:
        # Check if we should trigger (based on recent download)
        if session.get('trigger_power_automate'):
            # Clear the flag
            session['trigger_power_automate'] = False
            
            # Trigger in background
            threading.Thread(target=delayed_power_automate_trigger, args=(0,)).start()
            
            return jsonify({
                'status': 'success', 
                'message': 'Power Automate triggered automatically after ROB download'
            })
        else:
            return jsonify({
                'status': 'error', 
                'message': 'No recent ROB download detected'
            })
    except Exception as e:
        return jsonify({
            'status': 'error', 
            'message': f'Error: {str(e)}'
        })

# ============================================================================
# WEEKLY REPORT ROUTES
# ============================================================================

@app.route('/weekly_report', methods=['GET', 'POST'])
def weekly_report():
    if request.method == 'POST':
        form_type = request.form.get('form_type')
        
        if form_type == 'backend_processing':
            return handle_backend_processing()
        else:
            return handle_rid_analysis()
    
    # GET request - show form (no data to display)
    return render_template('weekly_report.html', qualified_rids=None, filter_summary=None, backend_result=None)
def handle_rid_analysis():
    """Handle RID analysis with dual file input - ranking sheet + cleaned ROB file"""
    try:
        print("RID Analysis POST request received!")
        
        # Get filter parameters from form
        min_search_volume = int(request.form.get('min_search_volume', 5000))
        competition_level = request.form.get('competition_level', 'Low')
        analyze_trends = request.form.get('analyze_trends') == 'on'
        
        print(f"User Filters: Search >= {min_search_volume}, Competition = {competition_level}")
        print(f"Google Trends: {'Enabled' if analyze_trends else 'Disabled'}")
        
        # Validate form inputs
        if not min_search_volume or min_search_volume < 0:
            flash('❌ Please enter a valid minimum search volume!')
            return redirect(request.url)
            
        if not competition_level:
            flash('❌ Please select a competition level!')
            return redirect(request.url)
        
        # Handle RANKING SHEET upload
        ranking_file = request.files.get('ranking_file')
        if not ranking_file or ranking_file.filename == '':
            flash('❌ Please select a ranking Excel file!')
            return redirect(request.url)

        if not allowed_file(ranking_file.filename):
            flash('❌ Only Excel files (.xlsx, .xls) and CSV files are allowed for ranking sheet!')
            return redirect(request.url)

        # Handle CLEANED ROB FILE upload
        rob_file = request.files.get('cleaned_rob_file')
        if not rob_file or rob_file.filename == '':
            flash('❌ Please select a cleaned ROB Excel file!')
            return redirect(request.url)

        if not allowed_file(rob_file.filename):
            flash('❌ Only Excel files (.xlsx, .xls) and CSV files are allowed for ROB file!')
            return redirect(request.url)

        # Save both uploaded files
        ranking_filename = secure_filename(ranking_file.filename)
        ranking_path = os.path.join(app.config['UPLOAD_FOLDER'], ranking_filename)
        ranking_file.save(ranking_path)
        print(f"Ranking file saved: {ranking_path}")
        
        rob_filename = secure_filename(rob_file.filename)
        rob_path = os.path.join(app.config['UPLOAD_FOLDER'], rob_filename)
        rob_file.save(rob_path)
        print(f"ROB file saved: {rob_path}")
        
        # Process both files and get qualified ROB data
        result_summary = process_dual_files_and_extract_rob(
            ranking_path, rob_path, min_search_volume, competition_level, analyze_trends
        )
        
        # Format success/warning messages based on results
        if result_summary['success']:
            flash(f'✅ Success! Found {result_summary["qualified_rids_count"]} qualified RIDs')
            flash(f'✅ Extracted {result_summary["matched_rob_rows"]} matching ROB rows')
            flash(f'📁 ROB.xlsx saved to Desktop/RPA folder!')
            print(f"Dual file processing completed: {result_summary}")
        else:
            flash(f'❌ Error: {result_summary.get("error", "Unknown error")}')
            result_summary = None
        
        # Clean up uploaded files after processing
        try:
            os.remove(ranking_path)
            os.remove(rob_path)
            print(f"Cleaned up uploaded files")
        except Exception as cleanup_error:
            print(f"Warning: Could not clean up files: {cleanup_error}")
        
        # Render template with results
        return render_template('weekly_report.html', 
                              qualified_rids=result_summary.get('qualified_rids', []) if result_summary else [],
                              filter_summary=result_summary.get('filter_summary', {}) if result_summary else {},
                              backend_result=None,
                              rob_extraction_result=result_summary)
        
    except ValueError as ve:
        print(f"Value Error: {ve}")
        flash('❌ Invalid input values. Please check your filters.')
        return redirect(request.url)
    except Exception as e:
        print(f"Error: {e}")
        flash(f'❌ Error processing files: {str(e)}')
        return redirect(request.url)


def process_dual_files_and_extract_rob(ranking_path, rob_path, min_search_volume, competition_level, analyze_trends):
    """Process ranking sheet and ROB file together, extract matching rows"""
    try:
        print(f"\n=== PROCESSING DUAL FILES ===")
        print(f"Ranking file: {ranking_path}")
        print(f"ROB file: {rob_path}")
        
        # STEP 1: Process ranking sheet to get qualified RIDs
        print("\n📊 STEP 1: Processing ranking sheet...")
        qualified_rids, filter_summary, updated_ranking_path = get_qualified_rids_and_remove_trending(
            ranking_path, min_search_volume, competition_level, analyze_trends
        )
        
        if not qualified_rids:
            return {
                'success': False,
                'error': 'No qualified RIDs found in ranking sheet with your filter criteria'
            }
        
        print(f"✅ Found {len(qualified_rids)} qualified RIDs: {qualified_rids}")
        
        # STEP 2: Process ROB file and extract matching rows
        print(f"\n📋 STEP 2: Processing ROB file and extracting matching rows...")
        
        # Read the cleaned ROB file
        if rob_path.endswith('.csv'):
            rob_df = pd.read_csv(rob_path)
        else:
            rob_df = pd.read_excel(rob_path, engine='openpyxl')
        
        total_rob_rows = len(rob_df)
        print(f"ROB file loaded: {total_rob_rows} rows")
        print(f"ROB file columns: {list(rob_df.columns)}")
        
        # Find Report ID column
        report_id_column = None
        possible_columns = ['Report ID', 'ReportID', 'report_id', 'ID', 'Report_ID', 'Market Name']
        
        for col in possible_columns:
            if col in rob_df.columns:
                report_id_column = col
                break
        
        if not report_id_column:
            return {
                'success': False,
                'error': f'Report ID column not found in ROB file. Available columns: {list(rob_df.columns)}'
            }
        
        print(f"Using Report ID column: {report_id_column}")
        
        # Convert qualified_rids to same type as Report ID column
        rob_df[report_id_column] = rob_df[report_id_column].astype(str).str.strip()
        qualified_rids_str = [str(rid).strip() for rid in qualified_rids]
        
        print(f"Sample Report IDs in ROB file: {rob_df[report_id_column].head().tolist()}")
        print(f"Looking for RIDs: {qualified_rids_str}")
        
        # Filter ROB rows that match qualified RIDs
        matching_rob_rows = rob_df[rob_df[report_id_column].isin(qualified_rids_str)].copy()
        matched_count = len(matching_rob_rows)
        
        print(f"Found {matched_count} matching ROB rows")
        
        if matched_count == 0:
            return {
                'success': False,
                'error': f'No matching Report IDs found in ROB file. Check if Report IDs {qualified_rids} exist in the ROB file.'
            }
        
        # Show which RIDs were found and missing
        found_rids = matching_rob_rows[report_id_column].tolist()
        missing_rids = [rid for rid in qualified_rids_str if rid not in found_rids]
        
        print(f"Found Report IDs: {found_rids}")
        if missing_rids:
            print(f"Missing Report IDs: {missing_rids}")
        
        # STEP 3: Save to Desktop/RPA folder as ROB.xlsx
        print(f"\n💾 STEP 3: Saving to Desktop...")
        
        # Create RPA folder on Desktop if it doesn't exist
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        rpa_folder = os.path.join(desktop_path, "RPA")
        
        if not os.path.exists(rpa_folder):
            os.makedirs(rpa_folder)
            print(f"Created RPA folder: {rpa_folder}")
        
        # Save as ROB.xlsx in Desktop/RPA folder
        output_path = os.path.join(rpa_folder, "weekly_ROB.xlsx")
        
        # Use xlsxwriter for better performance
        with pd.ExcelWriter(output_path, engine='xlsxwriter') as writer:
            matching_rob_rows.to_excel(writer, index=False, sheet_name='ROB_Data')
        
        print(f"✅ weekly_ROB.xlsx saved to: {output_path}")
        
        # Display sample of extracted data
        print("\nSample of extracted ROB data:")
        print(matching_rob_rows.head(2))
        
        # Create comprehensive summary
        summary = {
            'success': True,
            'qualified_rids': qualified_rids,
            'qualified_rids_count': len(qualified_rids),
            'total_rob_rows': total_rob_rows,
            'matched_rob_rows': matched_count,
            'found_rids': found_rids,
            'missing_rids': missing_rids,
            'output_path': output_path,
            'report_id_column': report_id_column,
            'filter_summary': filter_summary
        }
        
        return summary
        
    except Exception as e:
        print(f"Error in dual file processing: {e}")
        return {
            'success': False,
            'error': str(e)
        }

# 🔥 COMPLETE WORKFLOW SUMMARY:


def handle_backend_processing():
    """Handle backend file processing form submission"""
    try:
        print("Backend Processing POST request received!")
        
        # Get processing options
        auto_detect_header = request.form.get('auto_detect_header') == 'on'
        clean_columns = request.form.get('clean_columns') == 'on'
        remove_empty_rows = request.form.get('remove_empty_rows') == 'on'
        
        print(f"Processing options: Header={auto_detect_header}, Clean={clean_columns}, Remove Empty={remove_empty_rows}")
        
        # Handle backend file upload
        backend_file = request.files.get('backend_file')
        if not backend_file or backend_file.filename == '':
            flash('❌ Please select a backend Excel file!')
            return redirect(request.url)

        if not allowed_backend_file(backend_file.filename):
            flash('❌ Only Excel files (.xlsx, .xls) are allowed for backend processing!')
            return redirect(request.url)

        # Save uploaded backend file
        backend_filename = secure_filename(backend_file.filename)
        backend_path = os.path.join(app.config['UPLOAD_FOLDER'], backend_filename)
        backend_file.save(backend_path)
        print(f"Backend file saved: {backend_path}")
        
        # Process the backend file
        backend_result = process_backend_file(
            backend_path, 
            auto_detect_header=auto_detect_header, 
            clean_columns=clean_columns, 
            remove_empty_rows=remove_empty_rows
        )
        
        # Clean up uploaded file after processing
        try:
            os.remove(backend_path)
            print(f"Cleaned up backend file: {backend_path}")
        except Exception as cleanup_error:
            print(f"Warning: Could not clean up backend file {backend_path}: {cleanup_error}")
        
        # Format success/error messages
        if backend_result['success']:
            flash(f'✅ Backend file processed successfully!')
            flash(f'📁 Processed {backend_result["final_rows"]} rows from {backend_result["original_rows"]} original rows')
            flash(f'📥 ROB.xlsx ready for download!')
        else:
            flash(f'❌ Backend processing failed: {backend_result["error"]}')
        
        # Render template with backend results
        return render_template('weekly_report.html', 
                             qualified_rids=None,
                             filter_summary=None,
                             backend_result=backend_result)
        
    except Exception as e:
        print(f"Backend processing error: {e}")
        flash(f'❌ Error processing backend file: {str(e)}')
        return redirect(request.url)

def get_qualified_rids_and_remove_trending(file_path, min_search_volume, competition_level, analyze_trends=False):
    """Apply custom filters, get qualified RIDs, and remove trending rows from original sheet"""
    try:
        print(f"Processing file: {file_path}")
        
        # Read the file
        if file_path.endswith('.csv'):
            df = pd.read_csv(file_path)
        else:
            df = pd.read_excel(file_path, engine='openpyxl')
        
        original_count = len(df)
        print(f"Original data loaded: {original_count} rows")
        
        # Validate required columns exist
        required_columns = ['AVG. Search', 'Competition', 'RID']
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            raise ValueError(f"Missing required columns: {missing_columns}")
        
        # Apply user-defined filters
        print(f"Applying filters: AVG. Search >= {min_search_volume} AND Competition = {competition_level}")
        
        if competition_level == 'All':
            # No competition filter
            filtered_df = df[df['AVG. Search'] >= min_search_volume].copy()
        else:
            # Apply both search volume and competition filters
            filtered_df = df[
                (df['AVG. Search'] >= min_search_volume) & 
                (df['Competition'] == competition_level)
            ].copy()
        
        filtered_count = len(filtered_df)
        print(f"After applying filters: {filtered_count} rows")
        
        # Create filter summary
        filter_summary = {
            'min_search': f"{min_search_volume:,}",
            'competition': competition_level,
            'original_count': original_count,
            'filtered_count': filtered_count,
            'trends_enabled': analyze_trends
        }
        
        updated_file_path = None
        
        if filtered_count == 0:
            print("No records match the filter criteria")
            return [], filter_summary, updated_file_path
        
        if analyze_trends:
            # Run Google Trends analysis on filtered data
            print("🔥 Running Google Trends analysis on filtered keywords...")
            
            # Check if API key is configured
            if not GOOGLE_TRENDS_CONFIG.get('SERP_API_KEY') or GOOGLE_TRENDS_CONFIG['SERP_API_KEY'] == 'YOUR_SERP_API_KEY_HERE':
                print("⚠️ No SERP API key configured - returning all filtered RIDs")
                qualified_rids = filtered_df['RID'].tolist()
                return qualified_rids, filter_summary, updated_file_path
            
            # Run actual Google Trends analysis
            keywords_data = filtered_df.to_dict('records')
            trending_data = analyze_keywords_with_google_trends(keywords_data)
            qualified_rids = [item['RID'] for item in trending_data if 'RID' in item]
            
            print(f"Google Trends analysis complete: {len(qualified_rids)} trending RIDs out of {filtered_count} filtered")
            filter_summary['trends_qualified'] = len(qualified_rids)
            filter_summary['trends_message'] = f"After Google Trends analysis: {len(qualified_rids)} out of {filtered_count} keywords are trending"
            
            # Remove trending RIDs from original dataframe
            if qualified_rids:
                print(f"🗑️ Removing {len(qualified_rids)} trending RIDs from ranking sheet...")
                
                # Create a copy of original dataframe
                df_updated = df.copy()
                
                # Remove rows where RID is in the qualified_rids list
                df_updated = df_updated[~df_updated['RID'].isin(qualified_rids)]
                
                rows_removed = len(df) - len(df_updated)
                print(f"✅ Removed {rows_removed} trending rows from ranking sheet")
                
                # Save the updated ranking sheet
                updated_file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'updated_ranking_sheet.xlsx')
                
                # Use xlsxwriter for better performance
                with pd.ExcelWriter(updated_file_path, engine='xlsxwriter') as writer:
                    df_updated.to_excel(writer, index=False, sheet_name='Sheet1')
                
                print(f"💾 Updated ranking sheet saved as: {updated_file_path}")
                
                # Update filter summary with removal info
                filter_summary['rows_removed'] = rows_removed
                filter_summary['final_sheet_rows'] = len(df_updated)
                filter_summary['removal_message'] = f"Removed {rows_removed} trending rows. Updated sheet has {len(df_updated)} rows."
            else:
                print("ℹ️ No trending RIDs found - ranking sheet unchanged")
                filter_summary['removal_message'] = "No trending RIDs found - ranking sheet unchanged"
            
        else:
            # No Google Trends - return all filtered RIDs
            qualified_rids = filtered_df['RID'].tolist()
            print(f"Returning all filtered RIDs: {len(qualified_rids)} RIDs")
            filter_summary['trends_message'] = "Google Trends analysis disabled - showing all filtered results"
        
        return qualified_rids, filter_summary, updated_file_path
        
    except Exception as e:
        print(f"Error in get_qualified_rids_and_remove_trending: {e}")
        raise e

def process_backend_file(file_path, auto_detect_header=True, clean_columns=True, remove_empty_rows=True):
    """Process large backend file directly to ROB format with optimization"""
    try:
        print(f"\n=== PROCESSING LARGE BACKEND FILE TO ROB FORMAT ===")
        print(f"Processing file: {file_path}")
        
        # Step 1: Read the file with optimization for large files
        try:
            # Try reading with openpyxl engine for better large file handling
            df_raw = pd.read_excel(file_path, header=None, engine='openpyxl')
        except Exception as e:
            print(f"Error with openpyxl, trying alternative: {e}")
            # Fallback to default engine
            df_raw = pd.read_excel(file_path, header=None)
        
        original_rows = df_raw.shape[0]
        print(f"Initial raw data shape: {df_raw.shape}")
        
        # Step 2: Find the actual header row if auto-detect is enabled
        if auto_detect_header:
            header_row_index = find_header_row(df_raw)
        else:
            header_row_index = 0  # Assume first row is header
        
        if header_row_index is not None:
            # Set the header
            header = df_raw.iloc[header_row_index]
            # Drop rows before the header (inclusive)
            df_data = df_raw[header_row_index + 1:].copy()
            # Assign the correct header
            df_data.columns = header
            
            # Reset index
            df_data.reset_index(drop=True, inplace=True)
            
            print(f"Data extracted with header found at index {header_row_index}. New shape: {df_data.shape}")
            
            if clean_columns:
                # Clean column names (remove leading/trailing spaces, handle duplicates)
                df_data.columns = df_data.columns.str.strip()
                
                # Handle duplicate columns
                cols = pd.Series(df_data.columns)
                for dup in cols[cols.duplicated()].unique():
                    cols[cols[cols == dup].index.values.tolist()] = [dup + '_' + str(i) if i != 0 else dup for i in range(sum(cols == dup))]
                df_data.columns = cols
                
                print("Columns cleaned.")
            
            if remove_empty_rows:
                # Drop rows that are entirely null after extraction
                initial_rows = df_data.shape[0]
                df_data.dropna(how='all', inplace=True)
                rows_dropped = initial_rows - df_data.shape[0]
                print(f"Dropped {rows_dropped} empty rows.")
            
            # Save the processed file as "ROB.xlsx" directly
            output_path = os.path.join(app.config['UPLOAD_FOLDER'], 'ROB.xlsx')
            
            # Use xlsxwriter engine for better performance with large files
            with pd.ExcelWriter(output_path, engine='xlsxwriter') as writer:
                df_data.to_excel(writer, index=False, sheet_name='Sheet1')
            
            print(f"ROB file saved as: {output_path}")
            
            # Display sample of processed data
            print(f"\nProcessed {len(df_data)} records successfully")
            print("Sample of processed data:")
            print(df_data.head(2).to_string())
            
            # Create summary
            summary = {
                'success': True,
                'original_rows': original_rows,
                'final_rows': len(df_data),
                'header_row': header_row_index,
                'final_columns': len(df_data.columns),
                'output_file': 'ROB.xlsx'
            }
            
            return summary
        
        else:
            print("Could not automatically detect header row.")
            return {
                'success': False,
                'error': 'Could not automatically detect header row. Please check your file format.'
            }
    
    except MemoryError:
        print("Memory error - file too large")
        return {
            'success': False,
            'error': 'File too large to process. Please try with a smaller file or contact support.'
        }
    except Exception as e:
        print(f"Error processing backend file: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def find_header_row(df):
    """Heuristic function to find the header row"""
    for index, row in df.iterrows():
        if sum(isinstance(x, str) for x in row) >= 5:
            print(f"Potential header row found at index: {index}")
            return index
    return None

@app.route('/download_backend_file')
def download_backend_file():
    """Download the processed ROB file"""
    try:
        filename = 'ROB.xlsx'
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True, download_name=filename)
        else:
            flash('❌ Processed file not found. Please process a backend file first.')
            return redirect(url_for('weekly_report'))
    except Exception as e:
        flash(f'❌ Error downloading file: {str(e)}')
        return redirect(url_for('weekly_report'))

@app.route('/download_updated_ranking')
def download_updated_ranking():
    """Download the updated ranking sheet (with trending rows removed)"""
    try:
        filename = 'updated_ranking_sheet.xlsx'
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True, download_name='ranking_sheet_trending_removed.xlsx')
        else:
            flash('❌ Updated ranking sheet not found. Please run Google Trends analysis first.')
            return redirect(url_for('weekly_report'))
    except Exception as e:
        flash(f'❌ Error downloading updated ranking sheet: {str(e)}')
        return redirect(url_for('weekly_report'))

# ============================================================================
# GOOGLE TRENDS FUNCTIONS
# ============================================================================

def analyze_keywords_with_google_trends(keywords_data):
    """Analyze keywords with Google Trends"""
    api_key = GOOGLE_TRENDS_CONFIG['SERP_API_KEY']
    qualifying_keywords = []
    
    print(f"🔍 Analyzing {len(keywords_data)} filtered keywords with Google Trends...")
    
    for i, keyword_row in enumerate(keywords_data):
        try:
            original_keyword = keyword_row.get('Keywords', '')
            rid = keyword_row.get('RID', '')
            
            if not original_keyword or not rid:
                print(f"[{i+1}/{len(keywords_data)}] Skipping row with missing keyword or RID")
                continue
                
            clean_keyword = clean_keyword_for_trends(original_keyword)
            
            if not clean_keyword:
                print(f"[{i+1}/{len(keywords_data)}] Skipping empty keyword after cleaning: {original_keyword}")
                continue
            
            print(f"[{i+1}/{len(keywords_data)}] Analyzing RID {rid}: '{original_keyword}' → '{clean_keyword}'")
            
            # Check Google Trends conditions
            interest_condition = check_interest_over_time(clean_keyword, api_key)
            
            if interest_condition:
                if GOOGLE_TRENDS_CONFIG.get('REQUIRE_COUNTRIES_CHECK', True):
                    time.sleep(GOOGLE_TRENDS_CONFIG['REQUEST_DELAY'])
                    countries_condition = check_top_countries(clean_keyword, api_key)
                    time.sleep(GOOGLE_TRENDS_CONFIG['REQUEST_DELAY'])
                else:
                    countries_condition = True
                    print(f"  ℹ️  Countries check disabled - accepting based on interest only")
                
                if countries_condition:
                    qualifying_keywords.append({'RID': rid, 'keyword': original_keyword})
                    print(f"  ✅ TRENDING: RID {rid} - {original_keyword}")
                else:
                    print(f"  ❌ Not trending (countries): RID {rid} - {original_keyword}")
            else:
                print(f"  ❌ Not trending (interest): RID {rid} - {original_keyword}")
                
        except Exception as e:
            print(f"  ❌ Error analyzing RID {keyword_row.get('RID', 'unknown')}: {e}")
            continue
    
    return qualifying_keywords

def clean_keyword_for_trends(keyword):
    """Clean keyword by removing problematic terms"""
    if not keyword:
        return ""
        
    cleaned = str(keyword)
    
    # Remove terms from config
    for term in GOOGLE_TRENDS_CONFIG['TERMS_TO_REMOVE']:
        cleaned = re.sub(rf'\b{re.escape(term)}\b', '', cleaned, flags=re.IGNORECASE)
    
    # Clean up extra spaces and trim
    cleaned = ' '.join(cleaned.split()).strip()
    return cleaned

def check_interest_over_time(keyword, api_key):
    """Check if interest crosses threshold multiple times"""
    try:
        params = {
            'engine': 'google_trends',
            'q': keyword,
            'data_type': 'TIMESERIES',
            'date': 'now 7-d',
            'geo': '',
            'api_key': api_key
        }
        
        response = requests.get("https://serpapi.com/search", params=params, timeout=30)
        if response.status_code != 200:
            print(f"    API Error: {response.status_code}")
            return False
            
        data = response.json()
        
        if 'interest_over_time' in data and 'timeline_data' in data['interest_over_time']:
            timeline_data = data['interest_over_time']['timeline_data']
            crosses_threshold_count = 0
            
            for point in timeline_data:
                if 'values' in point:
                    for value in point['values']:
                        try:
                            interest_value = value.get('value', 0)
                            if interest_value is None:
                                interest_value = 0
                            elif isinstance(interest_value, str):
                                interest_value = interest_value.replace('<', '').replace('+', '').replace('>', '')
                                interest_value = float(interest_value) if interest_value.replace('.', '', 1).isdigit() else 0
                            else:
                                interest_value = float(interest_value)
                                
                            if interest_value >= GOOGLE_TRENDS_CONFIG['INTEREST_THRESHOLD']:
                                crosses_threshold_count += 1
                                break
                        except (ValueError, TypeError) as ve:
                            print(f"    Warning: Could not parse interest value: {value.get('value', 'N/A')} - {ve}")
                            continue
            
            print(f"    Interest: {crosses_threshold_count} days >= {GOOGLE_TRENDS_CONFIG['INTEREST_THRESHOLD']}")
            return crosses_threshold_count >= GOOGLE_TRENDS_CONFIG['DAYS_ABOVE_THRESHOLD']
        
        print("    No interest data available")
        return False
        
    except requests.exceptions.RequestException as e:
        print(f"    Network error checking interest for '{keyword}': {e}")
        return False
    except Exception as e:
        print(f"    Error checking interest for '{keyword}': {e}")
        return False

def check_top_countries(keyword, api_key):
    """Check if target countries appear in top regions"""
    try:
        params = {
            'engine': 'google_trends',
            'q': keyword,
            'data_type': 'GEO_MAP',
            'date': 'now 7-d',
            'api_key': api_key
        }
        
        response = requests.get("https://serpapi.com/search", params=params, timeout=30)
        
        if response.status_code != 200:
            return False
            
        data = response.json()
        
        if 'interest_by_region' in data:
            regions_data = data['interest_by_region']
            if not regions_data:
                return False
                
            top_regions = regions_data[:10]
            found_countries = []
            
            for region in top_regions:
                region_name = region.get('location', '').lower()
                
                for target_country in GOOGLE_TRENDS_CONFIG['TARGET_COUNTRIES']:
                    if target_country.lower() in region_name or region_name in target_country.lower():
                        found_countries.append(target_country)
                        break
            
            return len(found_countries) > 0
        
        return False
        
    except Exception as e:
        print(f"    Error checking countries for '{keyword}': {e}")
        return False

# ============================================================================
# POWER AUTOMATE ROUTES
# ============================================================================

@app.route('/wait_power_automate')
def wait_power_automate():
    """Show a waiting page for Power Automate Desktop step."""
    return render_template('wait_power_automate.html')

@app.route('/api/trigger_power_automate', methods=['POST'])
def trigger_power_automate_flow():
    """Triggers a Power Automate Desktop flow"""
    pad_exe_path = r"C:\Program Files (x86)\Power Automate Desktop\PAD.Console.Host.exe"
    flow_name = "Paid PR - Files Downloader"
    
    if not os.path.exists(pad_exe_path):
        print("Power Automate Desktop executable not found!")
        return jsonify({'status': 'error', 'message': 'PAD executable not found'})
    
    command = f'"{pad_exe_path}" -flow "{flow_name}"'
    
    try:
        result = subprocess.run(command, shell=True, check=True, text=True, capture_output=True)
        print(f"Flow triggered successfully. Output: {result.stdout}")

        time.sleep(5)
        
        flow_button_coordinates = (463, 395)
        print(f"Clicking at {flow_button_coordinates}")
        pyautogui.click(flow_button_coordinates)
        print("Flow triggered successfully.")

    except subprocess.CalledProcessError as e:
        print(f"Error triggering flow: {e.stderr}")
        return jsonify({'status': 'error', 'message': f'Flow error: {e.stderr}'})
    
    return jsonify({'status': 'success', 'message': 'Power Automate process completed.'})

# ============================================================================
# MAIN APPLICATION RUNNER
# ============================================================================

if __name__ == '__main__':
    import webbrowser
    webbrowser.open('http://127.0.0.1:5000/')
    app.run(debug=True, host='0.0.0.0', port=5000)